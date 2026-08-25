"""AI subtitle organization review workflow.

Layered model (SOURCE BASELINE is immutable):

- SOURCE BASELINE: the original Chirp/ASR segments in subtitles.json /
  subtitles-corrected.json. segment_id, start_ms, end_ms and raw_text are
  never modified or deleted by this module.

- CANDIDATE REVISION: AI-proposed changes stored as pending candidates.
  Accept/reject/edit only mutates candidate state; the active revision is
  untouched until an explicit publish.

- ACTIVE SUBTITLE REVISION: derived presentation cues. Every derived cue
  traces back to source_segment_ids and reuses source timestamps only.
  Publishing creates a new immutable Revision snapshot; rollback creates a
  new revision whose content equals the target, never a destructive
  overwrite.

Exports (SRT/VTT/TXT/DOCX) render from the active revision by default; any
historical revision can be rendered explicitly.

AI proposals are validated fail-closed:
- unknown/invented segment ids are rejected;
- invented timestamps are rejected (derived cues reuse source boundaries);
- cross-segment reflows must cover contiguous adjacent segments;
- segments with distinct speakers are never merged;
- every derived cue must carry a source lineage mapping;
- text conservation: reflow may redistribute characters across cues but the
  concatenated content must equal the scope input modulo whitespace, unless
  the change is a pure per-segment correction handled elsewhere.
"""
from __future__ import annotations

import hashlib
import json
import re
import threading
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from fastapi import APIRouter, HTTPException, Request
from pydantic import BaseModel, ConfigDict, Field

from app.api import _mutation_actor

DATA_DIR = Path(__import__("os").environ.get("COURSE_TRANSCRIPT_DATA_DIR", "/app/data"))
router = APIRouter(prefix="/api/v1/subtitles", tags=["subtitles"])
_LOCK = threading.RLock()


def _jobs_dir() -> Path:
    return DATA_DIR / "jobs"

CHANGE_TYPES = frozenset(
    {
        "asr_typo",
        "proper_noun",
        "semantic_asr_error",
        "punctuation",
        "line_break",
        "repetition_cleanup",
        "obvious_speech_correction",
        "cross_segment_reflow",
        "merge_adjacent",
        "split_for_readability",
        "mixed",
    }
)
RISKS = frozenset({"low", "medium", "high"})
MAX_LINES_PER_CUE = 2
MAX_CHARS_PER_LINE = 20


class CandidateProposal(BaseModel):
    model_config = ConfigDict(extra="forbid")
    change_type: str
    source_segment_ids: list[str] = Field(min_length=1, max_length=8)
    # before entries are cue-aware: each item carries either a per-segment
    # snapshot ({"segment_id": ..., "text": ...}) or, when the candidate
    # re-edits an Active merged/reflowed cue, the full working-cue snapshot
    # ({"source_segment_ids": [...], "text": ...}). The union shape keeps
    # immutable source evidence separate from the working representation.
    before: list[dict[str, Any]]
    after: list[dict[str, Any]] = Field(min_length=1, max_length=8)
    reason: str = Field(min_length=1, max_length=500)
    confidence: float = Field(ge=0.0, le=1.0)
    risk: str
    high_review_required: bool = True


class OrganizeRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")
    expected_revision: int = Field(ge=0)
    candidates: list[CandidateProposal] = Field(min_length=1, max_length=2000)


class CandidateDecision(BaseModel):
    model_config = ConfigDict(extra="forbid")
    change_id: str
    decision: str = Field(pattern="^(accept|reject)$")
    edited_after: list[dict[str, Any]] | None = None


class PublishRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")
    base_revision: int = Field(ge=0)


# ---------------------------------------------------------------------------
# baseline + storage helpers
# ---------------------------------------------------------------------------


def _iso() -> str:
    return datetime.now(UTC).isoformat()


def _safe_id(value: str) -> str:
    if not re.fullmatch(r"[A-Za-z0-9._-]{1,160}", value) or value in {".", ".."}:
        raise HTTPException(status_code=404, detail="Subtitle not found")
    return value


def _read_json(path: Path, default: Any = None) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (FileNotFoundError, OSError, json.JSONDecodeError):
        return default


def _atomic_json(path: Path, payload: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    temporary.replace(path)


def _directory(subtitle_id: str) -> Path:
    """Resolve a job subtitle directory (jobs only for the AI review flow)."""
    directory = _jobs_dir() / _safe_id(subtitle_id)
    if not directory.is_dir():
        raise HTTPException(status_code=404, detail="Subtitle not found")
    return directory


def baseline_segments(directory: Path) -> list[dict[str, Any]]:
    """Immutable SOURCE BASELINE read-only view (raw ASR evidence)."""
    for name in ("subtitles.json",):
        payload = _read_json(directory / name, {})
        items = payload.get("segments") if isinstance(payload, dict) else None
        if isinstance(items, list) and items:
            return [
                {
                    "segment_id": str(item["segment_id"]),
                    "start_ms": int(item["start_ms"]),
                    "end_ms": int(item["end_ms"]),
                    "raw_text": str(item.get("raw_text", item.get("text", ""))),
                    **(
                        {"speaker": str(item["speaker"])}
                        if item.get("speaker") is not None
                        else {}
                    ),
                }
                for item in items
                if isinstance(item, dict) and item.get("segment_id") is not None
            ]
    raise HTTPException(status_code=409, detail="Subtitle segments are not ready")


def working_segments(directory: Path) -> list[dict[str, Any]]:
    """Canonical current working subtitle: Gemini corrected + manual edits.

    Reuses the existing Subtitle Editor resolution (subtitles-corrected.json
    → corrected_text → subtitle-editor.json edits overlay → current_text) so
    there is exactly ONE canonical current-state definition. SOURCE evidence
    fields (segment_id/start_ms/end_ms/raw_text) stay immutable.
    """
    from app.subtitles import editor as subtitle_editor

    current, _ = subtitle_editor._current_segments(directory)
    return [
        {
            "segment_id": item["segment_id"],
            "start_ms": item["start_ms"],
            "end_ms": item["end_ms"],
            "raw_text": item["raw_text"],
            "working_text": item["current_text"],
            **({"speaker": speaker} if (speaker := _baseline_speaker(directory, item["segment_id"])) else {}),
        }
        for item in current
    ]


_SPEAKER_CACHE: dict[int, dict[str, str]] = {}


def _baseline_speaker(directory: Path, segment_id: str) -> str | None:
    cache_key = directory.stat().st_ino
    mapping = _SPEAKER_CACHE.get(cache_key)
    if mapping is None:
        payload = _read_json(directory / "subtitles.json", {}) or {}
        mapping = {}
        for item in payload.get("segments") or []:
            if isinstance(item, dict) and item.get("segment_id") is not None:
                if item.get("speaker") is not None:
                    mapping[str(item["segment_id"])] = str(item["speaker"])
        _SPEAKER_CACHE[cache_key] = mapping
    return mapping.get(segment_id)


def _review_state(directory: Path) -> dict[str, Any]:
    state = _read_json(directory / "ai-subtitle-review.json", {})
    if not isinstance(state, dict):
        state = {}
    return {
        "revision": int(state.get("revision", 0)),
        "candidates": state.get("candidates") if isinstance(state.get("candidates"), list) else [],
        "revisions": state.get("revisions") if isinstance(state.get("revisions"), list) else [],
        "active_revision": state.get("active_revision"),
        "updated_at": state.get("updated_at"),
    }


def _save_state(directory: Path, state: dict[str, Any]) -> None:
    state["updated_at"] = _iso()
    _atomic_json(directory / "ai-subtitle-review.json", state)


def _revision_by_number(state: dict[str, Any], number: int | None) -> dict[str, Any] | None:
    for item in state["revisions"]:
        if int(item["revision"]) == number:
            return item
    return None


# ---------------------------------------------------------------------------
# validation (fail-closed)
# ---------------------------------------------------------------------------


def _validate_candidate(
    candidate: CandidateProposal,
    baseline_index: dict[str, dict[str, Any]],
    working_cues: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    if candidate.change_type not in CHANGE_TYPES:
        raise HTTPException(status_code=422, detail=f"未知的修改類型: {candidate.change_type}")
    if candidate.risk not in RISKS:
        raise HTTPException(status_code=422, detail=f"未知的風險等級: {candidate.risk}")

    ids = candidate.source_segment_ids
    for segment_id in ids:
        if segment_id not in baseline_index:
            # AI invented a segment id that does not exist in the baseline.
            raise HTTPException(status_code=422, detail=f"發明不存在的 segment: {segment_id}")

    numbers = [baseline_index[s]["_index"] for s in ids]
    if numbers != sorted(numbers):
        raise HTTPException(status_code=422, detail="source_segment_ids 必須依原始順序")
    if len(set(numbers)) != len(numbers):
        raise HTTPException(status_code=422, detail="source_segment_ids 重複")

    before_by_id: dict[str, str] = {}
    cue_before: list[dict[str, Any]] = []
    representation: str | None = None
    seen_before_lineage: set[str] = set()
    cue_lineage_order: list[str] = []
    for item in candidate.before:
        has_cue = item.get("source_segment_ids") is not None
        has_segment = item.get("segment_id") is not None
        if has_cue == has_segment:
            raise HTTPException(status_code=422, detail="before 項目必須且只能提供 segment_id 或 source_segment_ids 其中一種")
        current_representation = "cue" if has_cue else "segment"
        if representation is None:
            representation = current_representation
        elif representation != current_representation:
            raise HTTPException(status_code=422, detail="before 不可混用 per-segment 與 cue-aware snapshot")
        if has_cue:
            raw_ids = item.get("source_segment_ids")
            if not isinstance(raw_ids, list) or not raw_ids:
                raise HTTPException(status_code=422, detail="cue-aware before 必須有 source_segment_ids")
            ids_key = [str(v) for v in raw_ids]
            if len(set(ids_key)) != len(ids_key):
                raise HTTPException(status_code=422, detail="cue-aware before lineage 重複")
            for segment_id in ids_key:
                if segment_id in seen_before_lineage:
                    raise HTTPException(status_code=422, detail=f"before lineage 重複覆蓋: {segment_id}")
                seen_before_lineage.add(segment_id)
                cue_lineage_order.append(segment_id)
            cue_before.append({"source_segment_ids": ids_key, "text": str(item.get("text", ""))})
        else:
            key = str(item["segment_id"])
            if key in before_by_id:
                raise HTTPException(status_code=422, detail=f"before segment 重複: {key}")
            before_by_id[key] = str(item.get("text", ""))
    covered_before = set(before_by_id) | seen_before_lineage
    if covered_before != set(ids):
        raise HTTPException(status_code=422, detail="before 與 source_segment_ids 不一致")
    if cue_before and cue_lineage_order != list(ids):
        raise HTTPException(status_code=422, detail="cue-aware before 必須依 source_segment_ids 原始順序完整分割")
    if working_cues:
        scope = set(ids)
        for working_cue in working_cues:
            members = {str(v) for v in working_cue.get("source_segment_ids", [])}
            if len(members) > 1 and scope & members and not members <= scope:
                raise HTTPException(status_code=409, detail=f"候選僅涵蓋 Active multi-source cue 的部分 lineage；必須完整包含 {sorted(members)}")

    if cue_before:
        # Validate the cue snapshots against the canonical Active cues: the
        # candidate must re-edit the complete lineage with its current text.
        active_cues = {
            tuple(str(v) for v in cue.get("source_segment_ids", [])): str(cue.get("text", ""))
            for cue in (working_cues or [])
        }
        for snapshot in cue_before:
            key = tuple(snapshot["source_segment_ids"])
            expected = active_cues.get(key)
            if expected is None:
                raise HTTPException(
                    status_code=409,
                    detail=f"before 與目前 Active cue 不一致（lineage {list(key)} 不存在或已改變）",
                )
            if snapshot["text"] != expected:
                raise HTTPException(
                    status_code=422,
                    detail=f"before 文字與目前 Active cue 不一致: {list(key)}",
                )
    else:
        for segment_id, text in before_by_id.items():
            if working_cues:
                # Canonical truth is the Active Revision: the expected text
                # for a segment is the working cue that owns it.
                owning = [
                    str(cue.get("text", ""))
                    for cue in working_cues
                    if segment_id in [str(v) for v in cue.get("source_segment_ids", [])]
                ]
                expected = owning[0] if owning else str(
                    baseline_index[segment_id].get("working_text", baseline_index[segment_id]["raw_text"])
                )
            else:
                expected = str(baseline_index[segment_id].get("working_text", baseline_index[segment_id]["raw_text"]))
            if text != expected:
                raise HTTPException(status_code=422, detail=f"before 文字與來源不一致: {segment_id}")

    # Contiguity: ANY multi-segment candidate must span adjacent segments.
    indices = [baseline_index[s]["_index"] for s in ids]
    if len(ids) > 1 and max(indices) - min(indices) != len(indices) - 1:
        raise HTTPException(status_code=422, detail="跨段整理必須是相鄰 segments")

    seen_lineage: set[str] = set()
    total_after_chars = 0
    for cue in candidate.after:
        cue_ids = cue.get("source_segment_ids")
        if not isinstance(cue_ids, list) or not cue_ids:
            raise HTTPException(status_code=422, detail="每個 cue 必須有 source_segment_ids")
        for segment_id in cue_ids:
            if segment_id not in baseline_index:
                raise HTTPException(status_code=422, detail=f"cue 引用不存在的 segment: {segment_id}")
            if segment_id not in ids:
                raise HTTPException(status_code=422, detail=f"cue 超出候選範圍: {segment_id}")
            if segment_id in seen_lineage:
                raise HTTPException(
                    status_code=422,
                    detail=f"同一 source segment 不可出現在多個 derived cues: {segment_id}",
                )
            seen_lineage.add(segment_id)
        text = cue.get("text")
        if not isinstance(text, str) or not text.strip():
            raise HTTPException(status_code=422, detail="cue 文字不可空白")
        total_after_chars += len(text.strip())
        # No invented timestamps allowed anywhere in the payload.
        if "start_ms" in cue or "end_ms" in cue:
            raise HTTPException(status_code=422, detail="禁止發明 timestamp；時間僅能由來源 boundary 推導")
        # Speaker boundary (shared rule): a derived cue must never span
        # segments with different non-null speakers, regardless of type.
        speakers = {
            baseline_index[value].get("speaker")
            for value in cue_ids
            if baseline_index[value].get("speaker")
        }
        if len(speakers) > 1:
            raise HTTPException(status_code=422, detail="不同 speaker 的字幕不可合併為同一 derived cue")

    if seen_lineage != set(ids):
        missing = sorted(set(ids) - seen_lineage)
        raise HTTPException(status_code=422, detail=f"缺少來源對應的 segment: {missing}")

    multi = len(ids) > 1
    if multi and candidate.change_type not in {
        "cross_segment_reflow", "merge_adjacent", "split_for_readability", "mixed",
    }:
        raise HTTPException(status_code=422, detail=f"{candidate.change_type} 僅允許單一 segment")

    # Reflow/split must span contiguous adjacent segments only.
    if multi and candidate.change_type in {"cross_segment_reflow", "split_for_readability"}:
        indices = [baseline_index[s]["_index"] for s in ids]
        if max(indices) - min(indices) != len(indices) - 1:
            raise HTTPException(status_code=422, detail="跨段整理必須是相鄰 segments")

    if candidate.change_type == "merge_adjacent":
        if len(candidate.after) != 1:
            raise HTTPException(status_code=422, detail="merge_adjacent 只能產生單一 cue")
        if len({baseline_index[s].get("speaker") for s in ids if baseline_index[s].get("speaker")}) > 1:
            raise HTTPException(status_code=422, detail="不同 speaker 的字幕不可合併")

    # Pure reflow/merge must conserve visible text even for cue-aware R2 edits.
    if candidate.change_type in {"cross_segment_reflow", "merge_adjacent"}:
        before_joined = (
            "".join(str(snapshot["text"]) for snapshot in cue_before)
            if cue_before
            else "".join(before_by_id[s] for s in ids)
        )
        after_joined = "".join(str(cue["text"]) for cue in candidate.after)
        strip = lambda value: re.sub(r"\s+", "", value)  # noqa: E731
        if strip(before_joined) != strip(after_joined):
            raise HTTPException(
                status_code=422,
                detail="跨段整理不可改變總字詞內容；文字修正請使用 correction/mixed candidate",
            )

    # split_for_readability v1: presentation line wrapping only — one source
    # segment must never become multiple timed cues without word-level timing.
    if candidate.change_type == "split_for_readability" and len(candidate.after) > 1:
        raise HTTPException(
            status_code=422,
            detail="第一版不支援將單一 segment 拆成多個時間 cue；請使用單 cue 多行（line wrapping）",
        )
    for cue in candidate.after:
        cue_ids = [str(value) for value in cue["source_segment_ids"]]
        cue_indices = [baseline_index[value]["_index"] for value in cue_ids]
        if len(cue_indices) > 1 and max(cue_indices) - min(cue_indices) != len(cue_indices) - 1:
            raise HTTPException(status_code=422, detail="cue 來源必須連續")
        if sorted(cue_indices) != cue_indices:
            raise HTTPException(status_code=422, detail="cue source_segment_ids 必須依原始順序")

    return {
        "change_id": f"cand-{uuid.uuid4().hex[:12]}",
        "change_type": candidate.change_type,
        "source_segment_ids": list(ids),
        "before": cue_before if cue_before else [
            {"segment_id": segment_id, "text": before_by_id[segment_id]} for segment_id in ids
        ],
        "after": [
            {"source_segment_ids": list(cue["source_segment_ids"]), "text": str(cue["text"])}
            for cue in candidate.after
        ],
        "reason": candidate.reason,
        "confidence": candidate.confidence,
        "risk": candidate.risk,
        "high_review_required": True,
        "status": "pending",
        "decision_note": None,
        "created_at": _iso(),
    }


def _resolve_cues(
    directory: Path,
    state: dict[str, Any],
    accepted: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Materialize derived cues for Revision N+1.

    Invariant: Revision N+1 = Active Revision N + accepted candidate delta.
    Only when no active revision exists is the base the canonical current
    working subtitle (Gemini corrected + manual Subtitle Editor edits).
    Timestamps always come from source boundaries — never invented.
    """
    if state.get("active_revision") is not None:
        active = _revision_by_number(state, int(state["active_revision"]))
        if active is None:
            raise HTTPException(status_code=409, detail="Active revision 遺失")
        base_cues = [dict(cue) for cue in active["cues"]]
    else:
        base_cues = [
            {
                "text": str(item["working_text"]),
                "source_segment_ids": [item["segment_id"]],
                "start_ms": item["start_ms"],
                "end_ms": item["end_ms"],
            }
            for item in working_segments(directory)
        ]

    # Fail closed on overlapping accepted candidate scopes: two candidates
    # claiming the same source segment must never both apply (409, not guess).
    claimed_by: dict[str, str] = {}
    for candidate in accepted:
        for segment_id in candidate["source_segment_ids"]:
            previous = claimed_by.get(segment_id)
            if previous is not None and previous != candidate["change_id"]:
                raise HTTPException(
                    status_code=409,
                    detail=f"多個已接受候選同時修改 {segment_id}；請先解決衝突再發布",
                )
            claimed_by[segment_id] = candidate["change_id"]

    # V1 cannot safely edit only part of an already merged/reflowed Active cue.
    # A candidate that touches a multi-source Active cue must own its complete
    # lineage, otherwise replacing one member would either drop the remaining
    # members or trigger an internal lookup failure. Fail closed instead.
    for candidate in accepted:
        scope = {str(value) for value in candidate["source_segment_ids"]}
        for base_cue in base_cues:
            members = {str(value) for value in base_cue["source_segment_ids"]}
            if scope & members and not members <= scope:
                raise HTTPException(
                    status_code=409,
                    detail=(
                        "候選僅涵蓋 Active multi-source cue 的部分 lineage；"
                        f"必須完整包含 {sorted(members)}"
                    ),
                )

    replaced_segments: dict[str, list[dict[str, Any]]] = {}
    for candidate in accepted:
        cue_specs: list[dict[str, Any]] = []
        for cue in candidate["after"]:
            cue_ids = [str(value) for value in cue["source_segment_ids"]]
            timings = [_timing_of(base_cues, segment_id) or _baseline_timing(directory, segment_id) for segment_id in cue_ids]
            start = min(item[0] for item in timings)
            end = max(item[1] for item in timings)
            cue_specs.append(
                {
                    "text": str(cue["text"]),
                    "source_segment_ids": cue_ids,
                    "start_ms": start,
                    "end_ms": end,
                }
            )
        for segment_id in candidate["source_segment_ids"]:
            replaced_segments[segment_id] = cue_specs

    cues: list[dict[str, Any]] = []
    seen_replaced: set[str] = set()
    for cue in sorted(base_cues, key=lambda item: item["start_ms"]):
        members = [str(value) for value in cue["source_segment_ids"]]
        if any(member in replaced_segments for member in members):
            for member in members:
                if member not in seen_replaced:
                    seen_replaced.add(member)
                    specs = replaced_segments[member]
                    if all(spec not in cues for spec in specs):
                        cues.extend(specs)
            continue
        cues.append(dict(cue))
    cues.sort(key=lambda item: (item["start_ms"], item["end_ms"]))
    _validate_revision_cues(directory, cues)
    for index, cue in enumerate(cues, 1):
        cue["cue_id"] = f"cue-{index:04d}"
    return cues


def _validate_final_cues(cues: list[dict[str, Any]]) -> None:
    """Fail closed if a materialized revision has invalid timeline/lineage."""
    previous_end: int | None = None
    seen_lineage: set[str] = set()
    for cue in cues:
        start = int(cue["start_ms"])
        end = int(cue["end_ms"])
        if start >= end:
            raise HTTPException(status_code=409, detail="字幕 cue 時間範圍無效")
        if previous_end is not None and start < previous_end:
            raise HTTPException(status_code=409, detail="字幕 cue 時間重疊，拒絕建立 Revision")
        lineage = [str(value) for value in cue.get("source_segment_ids", [])]
        if not lineage:
            raise HTTPException(status_code=409, detail="字幕 cue 缺少 source lineage")
        for segment_id in lineage:
            if segment_id in seen_lineage:
                raise HTTPException(
                    status_code=409,
                    detail=f"Final Revision 重複使用 source segment: {segment_id}",
                )
            seen_lineage.add(segment_id)
        previous_end = end


def _validate_revision_cues(directory: Path, cues: list[dict[str, Any]]) -> None:
    """Validate timeline, unique lineage, and complete immutable-source coverage."""
    _validate_final_cues(cues)
    expected = {item["segment_id"] for item in baseline_segments(directory)}
    actual = {str(segment_id) for cue in cues for segment_id in cue.get("source_segment_ids", [])}
    if actual != expected:
        raise HTTPException(
            status_code=409,
            detail=f"Final Revision source coverage 不完整；missing={sorted(expected - actual)}, extra={sorted(actual - expected)}",
        )


def _timing_of(cues: list[dict[str, Any]], segment_id: str) -> tuple[int, int] | None:
    for cue in cues:
        if segment_id in [str(value) for value in cue["source_segment_ids"]]:
            return int(cue["start_ms"]), int(cue["end_ms"])
    return None


def _baseline_timing(directory: Path, segment_id: str) -> tuple[int, int]:
    for item in baseline_segments(directory):
        if item["segment_id"] == segment_id:
            return item["start_ms"], item["end_ms"]
    raise HTTPException(status_code=422, detail=f"未知 segment: {segment_id}")


# ---------------------------------------------------------------------------
# rendering
# ---------------------------------------------------------------------------


def _srt_time(value: int, separator: str = ",") -> str:
    value = max(0, int(value))
    hours, value = divmod(value, 3_600_000)
    minutes, value = divmod(value, 60_000)
    seconds, milliseconds = divmod(value, 1_000)
    return f"{hours:02}:{minutes:02}:{seconds:02}{separator}{milliseconds:03}"


def render_srt(cues: list[dict[str, Any]]) -> str:
    blocks = []
    for index, cue in enumerate(cues, 1):
        lines = _wrap_lines(cue["text"])
        blocks.append(
            f"{index}\n{_srt_time(cue['start_ms'])} --> {_srt_time(cue['end_ms'])}\n" + "\n".join(lines)
        )
    return "\n\n".join(blocks) + "\n"


def render_vtt(cues: list[dict[str, Any]]) -> str:
    body = ["WEBVTT", ""]
    for index, cue in enumerate(cues, 1):
        lines = _wrap_lines(cue["text"])
        body.append(str(index))
        body.append(f"{_srt_time(cue['start_ms'], '.')} --> {_srt_time(cue['end_ms'], '.')}")
        body.extend(lines)
        body.append("")
    return "\n".join(body)


def render_txt(cues: list[dict[str, Any]]) -> str:
    # Normalize whitespace, never remove it: "machine learning" must keep
    # its inner space while CJK text stays unwrapped.
    return "\n".join(re.sub(r"\s+", " ", cue["text"]).strip() for cue in cues) + "\n"


def _wrap_lines(text: str) -> list[str]:
    text = text.replace("\n", " ").strip()
    if len(text) <= MAX_CHARS_PER_LINE:
        return [text]
    midpoint = (len(text) + 1) // 2
    for pivot in range(midpoint, max(0, midpoint - 6), -1):
        if pivot < len(text) and text[pivot - 1] in "，。！？、；：":
            return [text[:pivot], text[pivot:]]
    return [text[:midpoint], text[midpoint:]]


def render_docx_bytes(cues: list[dict[str, Any]]) -> bytes:
    try:
        from docx import Document
        from io import BytesIO
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise HTTPException(status_code=503, detail="DOCX 匯出元件未安裝") from exc
    document = Document()
    document.add_heading("Course Transcript", level=0)
    buffer = BytesIO()
    for cue in cues:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"[{_srt_time(int(cue['start_ms']), ':')[:-4]}] ")
        run.bold = True
        paragraph.add_run(str(cue["text"]))
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# routes
# ---------------------------------------------------------------------------


@router.get("/{subtitle_id}/ai-review")
def get_ai_review_state(subtitle_id: str) -> dict[str, Any]:
    directory = _directory(subtitle_id)
    state = _review_state(directory)
    counts: dict[str, int] = {}
    for candidate in state["candidates"]:
        counts[candidate["status"]] = counts.get(candidate["status"], 0) + 1
    return {
        "subtitle_id": subtitle_id,
        "revision": state["revision"],
        "active_revision": state["active_revision"],
        "counts": counts,
        "total_candidates": len(state["candidates"]),
        "candidates": state["candidates"],
        "revisions": [
            {
                "revision": item["revision"],
                "created_at": item["created_at"],
                "source": item["source"],
                "content_sha256": item["content_sha256"],
                "cue_count": len(item["cues"]),
                "active": state["active_revision"] == item["revision"],
            }
            for item in state["revisions"]
        ],
    }


@router.post("/{subtitle_id}/ai-review/candidates")
def propose_candidates(subtitle_id: str, payload: OrganizeRequest, request: Request) -> dict[str, Any]:
    actor = _mutation_actor(request)
    directory = _directory(subtitle_id)
    with _LOCK:
        state = _review_state(directory)
        if state["revision"] != payload.expected_revision:
            raise HTTPException(status_code=409, detail="審核狀態已更新，請重新載入")
        working = working_segments(directory)
        baseline_index = {
            item["segment_id"]: {**item, "_index": position}
            for position, item in enumerate(working)
        }
        # Candidate validation input must equal the true canonical state:
        # Active Revision cues when published, otherwise editor working text.
        from app.subtitles import canonical_state

        cues, _ = canonical_state.canonical_cues(directory)
        base_identity = canonical_state.publication_identity(directory)
        created = []
        for candidate in payload.candidates:
            record = _validate_candidate(candidate, baseline_index, working_cues=cues)
            record["proposed_by"] = actor
            # Bind this review round to the canonical base it was generated
            # against; decide/publish re-check before applying.
            record["base_publication_key"] = base_identity["publication_key"]
            state["candidates"].append(record)
            created.append(record)
        _save_state(directory, state)
        return {
            "created": len(created),
            "total_candidates": len(state["candidates"]),
            "expected_revision": state["revision"],
            "candidates": created,
        }


@router.post("/{subtitle_id}/ai-review/candidates/decide")
def decide_candidate(subtitle_id: str, payload: CandidateDecision, request: Request) -> dict[str, Any]:
    actor = _mutation_actor(request)
    directory = _directory(subtitle_id)
    with _LOCK:
        state = _review_state(directory)
        target = next((item for item in state["candidates"] if item["change_id"] == payload.change_id), None)
        if target is None:
            raise HTTPException(status_code=404, detail="Candidate not found")
        if target["status"] != "pending":
            raise HTTPException(status_code=409, detail="此建議已審核")
        from app.subtitles import canonical_state

        current_key = canonical_state.publication_identity(directory)["publication_key"]
        if target.get("base_publication_key") not in (None, current_key):
            raise HTTPException(
                status_code=409,
                detail=(
                    "Canonical base 已改變；請清空本輪候選並以最新字幕重新產生提案 "
                    f"(base={target['base_publication_key']}, current={current_key})"
                ),
            )
        after = target["after"]
        if payload.decision == "accept":
            if payload.edited_after is not None:
                replacement = CandidateProposal(
                    change_type=target["change_type"],
                    source_segment_ids=target["source_segment_ids"],
                    before=target["before"],
                    after=payload.edited_after,
                    reason=target["reason"],
                    confidence=target["confidence"],
                    risk=target["risk"],
                    high_review_required=True,
                )
                baseline = working_segments(directory)
                baseline_index = {
                    item["segment_id"]: {**item, "_index": position}
                    for position, item in enumerate(baseline)
                }
                from app.subtitles import canonical_state

                cues_now, _ = canonical_state.canonical_cues(directory)
                validated = _validate_candidate(replacement, baseline_index, working_cues=cues_now)
                # Persist the normalized/validated edit so Publish uses the
                # human-approved text, not the original AI proposal. Keep the
                # AI original as audit metadata.
                if "original_after" not in target:
                    target["original_after"] = target["after"]
                target["after"] = validated["after"]
                target["manually_edited"] = True
                target["edited_at"] = _iso()
            target["status"] = "accepted"
        else:
            target["status"] = "rejected"
            if payload.edited_after is not None:
                raise HTTPException(status_code=422, detail="拒絕不可附帶修改內容")
        target["decided_by"] = actor
        target["decided_at"] = _iso()
        _save_state(directory, state)
        counts: dict[str, int] = {}
        for candidate in state["candidates"]:
            counts[candidate["status"]] = counts.get(candidate["status"], 0) + 1
        return {"change_id": payload.change_id, "status": target["status"], "counts": counts}


@router.post("/{subtitle_id}/ai-review/publish")
def publish_revision(subtitle_id: str, payload: PublishRequest, request: Request) -> dict[str, Any]:
    actor = _mutation_actor(request)
    directory = _directory(subtitle_id)
    with _LOCK:
        state = _review_state(directory)
        if state["revision"] != payload.base_revision:
            raise HTTPException(status_code=409, detail="審核狀態已更新，請重新載入")
        from app.subtitles import canonical_state

        current_key = canonical_state.publication_identity(directory)["publication_key"]
        stale_bases = {
            item.get("base_publication_key")
            for item in state["candidates"]
            if item.get("base_publication_key") not in (None, current_key)
        }
        if stale_bases:
            raise HTTPException(
                status_code=409,
                detail=(
                    "本輪候選的 canonical base 已改變，拒絕發布；請重新產生提案 "
                    f"(stale={sorted(str(k) for k in stale_bases)}, current={current_key})"
                ),
            )
        pending = sum(1 for item in state["candidates"] if item["status"] == "pending")
        if pending:
            raise HTTPException(status_code=409, detail=f"尚有 {pending} 項建議未審核")
        accepted = [item for item in state["candidates"] if item["status"] == "accepted"]
        cues = _resolve_cues(directory, state, accepted)
        digest = hashlib.sha256(render_srt(cues).encode("utf-8")).hexdigest()
        # publication_key must use the canonical identity hash so keys are
        # consistent across publish/status/cue-edit paths.
        from app.subtitles import canonical_state

        digest = canonical_state._cues_content_sha256(cues)
        number = state["revision"] + 1
        revision_record = {
            "revision": number,
            "created_at": _iso(),
            "created_by": actor,
            "source": "ai_subtitle_review_publish",
            "accepted_change_ids": [item["change_id"] for item in accepted],
            "rejected_change_ids": [
                item["change_id"] for item in state["candidates"] if item["status"] == "rejected"
            ],
            # Immutable audit snapshot: full candidate details must remain
            # recoverable after publish (not just IDs).
            "candidates_snapshot": [json.loads(json.dumps(item, ensure_ascii=False)) for item in state["candidates"]],
            "cues": cues,
            "content_sha256": digest,
        }
        state["revisions"].append(revision_record)
        previous_active = state["active_revision"]
        state["active_revision"] = number
        state["revision"] = number
        # Reviewed decisions are folded into the published revision; keep the
        # audit trail inside the revision record itself.
        state["candidates"] = []
        _save_state(directory, state)
        return {
            "revision": number,
            "previous_active_revision": previous_active,
            "active_revision": state["active_revision"],
            "cue_count": len(cues),
            "accepted_count": len(accepted),
            "content_sha256": digest,
            "publication_key": f"ai_review_active:r{number}:{digest[:16]}",
        }


@router.post("/{subtitle_id}/ai-review/revisions/{revision}/rollback")
def rollback_revision(subtitle_id: str, revision: int, request: Request) -> dict[str, Any]:
    actor = _mutation_actor(request)
    directory = _directory(subtitle_id)
    with _LOCK:
        state = _review_state(directory)
        target = _revision_by_number(state, revision)
        if target is None:
            raise HTTPException(status_code=404, detail="Revision not found")
        if revision == state["active_revision"]:
            raise HTTPException(status_code=409, detail="此版本已是 Active")
        if state["candidates"]:
            raise HTTPException(status_code=409, detail="尚有本輪 Review candidates；發布或清空後才能 rollback")
        cues = [dict(cue) for cue in target["cues"]]
        _validate_revision_cues(directory, cues)
        number = state["revision"] + 1
        record = {
            "revision": number,
            "created_at": _iso(),
            "created_by": actor,
            "source": "rollback",
            "rolled_back_from": revision,
            "accepted_change_ids": [],
            "rejected_change_ids": [],
            "cues": cues,
            "content_sha256": target["content_sha256"],
        }
        state["revisions"].append(record)
        state["active_revision"] = number
        state["revision"] = number
        _save_state(directory, state)
        return {"revision": number, "restored_from": revision, "active_revision": number}


@router.get("/{subtitle_id}/ai-review/export/{kind}")
@router.get("/{subtitle_id}/ai-review/export/{kind}/{revision}")
def export_revision(subtitle_id: str, kind: str, revision: int | None = None) -> Any:
    from fastapi import Response

    directory = _directory(subtitle_id)
    state = _review_state(directory)
    resolved = revision if revision is not None else state["active_revision"]
    if resolved is None:
        raise HTTPException(status_code=409, detail="尚無已發布的字幕版本")
    record = _revision_by_number(state, resolved)
    if record is None:
        raise HTTPException(status_code=404, detail="Revision not found")
    cues = [dict(cue) for cue in record["cues"]]
    _validate_revision_cues(directory, cues)
    kind = kind.lower()
    if kind == "srt":
        return Response(render_srt(cues), media_type="application/x-subrip")
    if kind == "vtt":
        return Response(render_vtt(cues), media_type="text/vtt")
    if kind == "txt":
        return Response(render_txt(cues), media_type="text/plain")
    if kind == "docx":
        return Response(
            render_docx_bytes(cues),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{subtitle_id}-r{resolved}.docx"'},
        )
    raise HTTPException(status_code=422, detail="不支援的匯出格式")


def _canonical_cue_hash(cues: list[dict[str, Any]]) -> str:
    """Canonical content hash shared with canonical_state.publication_identity."""
    from app.subtitles import canonical_state

    return canonical_state._cues_content_sha256(cues)


class CueEditRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")    # The cue being edited, identified by full lineage + current text
    # (optimistic concurrency against the Active revision).
    source_segment_ids: list[str] = Field(min_length=1, max_length=8)
    current_text: str = Field(min_length=1)
    new_text: str = Field(min_length=1)
    base_content_sha256: str | None = None


@router.patch("/{subtitle_id}/ai-review/cue")
def edit_canonical_cue(subtitle_id: str, payload: CueEditRequest, request: Request) -> dict[str, Any]:
    """Cue-aware human edit of the canonical current subtitle.

    Editing a canonical cue's text (structure unchanged) creates an auditable
    Revision N+1 from Active N; historical revisions stay immutable. Splitting
    a merged cue without word-level immutable timing is fail-closed.
    """
    actor = _mutation_actor(request)
    directory = _directory(subtitle_id)
    with _LOCK:
        from app.subtitles import canonical_state

        cues, source = canonical_state.canonical_cues(directory)
        if source != "ai_review_active":
            raise HTTPException(
                status_code=409,
                detail="目前無 AI Active Revision；請使用既有 per-segment Editor 流程",
            )
        state = _review_state(directory)
        if payload.base_content_sha256 is not None:
            identity = canonical_state.publication_identity(directory)
            if payload.base_content_sha256 != identity["content_sha256"]:
                raise HTTPException(
                    status_code=409,
                    detail="Canonical content changed since load；請重新載入",
                )
        target = next(
            (
                cue
                for cue in cues
                if [str(v) for v in cue["source_segment_ids"]]
                == [str(v) for v in payload.source_segment_ids]
            ),
            None,
        )
        if target is None:
            raise HTTPException(status_code=404, detail="找不到對應的 canonical cue")
        if str(target["text"]) != payload.current_text:
            raise HTTPException(status_code=409, detail="cue 文字已改變，請重新載入")

        new_cues = []
        for cue in cues:
            if cue is target:
                lineage = [str(v) for v in cue["source_segment_ids"]]
                if len(lineage) > 0:
                    new_cues.append({**cue, "text": payload.new_text.strip()})
                else:  # pragma: no cover - guarded by min_length
                    raise HTTPException(status_code=422, detail="lineage 不可為空")
            else:
                new_cues.append(dict(cue))
        _validate_revision_cues(directory, new_cues)

        number = state["revision"] + 1
        record = {
            "revision": number,
            "created_at": _iso(),
            "created_by": actor,
            "source": "editor_cue_edit",
            "edited_lineage": [str(v) for v in payload.source_segment_ids],
            "accepted_change_ids": [],
            "rejected_change_ids": [],
            "cues": [
                {
                    "text": str(cue["text"]),
                    "source_segment_ids": [str(v) for v in cue["source_segment_ids"]],
                    "start_ms": int(cue["start_ms"]),
                    "end_ms": int(cue["end_ms"]),
                }
                for cue in new_cues
            ],
            "content_sha256": _canonical_cue_hash(new_cues),
        }
        state["revisions"].append(record)
        state["active_revision"] = number
        state["revision"] = number
        _save_state(directory, state)
        return {
            "revision": number,
            "active_revision": number,
            "publication_key": (
                f"ai_review_active:r{number}:{record['content_sha256'][:16]}"
            ),
            "edited_lineage": [str(v) for v in payload.source_segment_ids],
        }


class BatchCueReplaceRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")
    replacements: list[dict[str, Any]] = Field(min_length=1, max_length=5000)
    base_content_sha256: str | None = None


@router.post("/{subtitle_id}/ai-review/batch-replace-cues")
def batch_replace_canonical_cues(
    subtitle_id: str, payload: BatchCueReplaceRequest, request: Request
) -> dict[str, Any]:
    """Batch replace against canonical working cues → auditable new revision.

    All edits apply atomically as one new Revision; any mismatch fails
    closed without partial application.
    """
    actor = _mutation_actor(request)
    directory = _directory(subtitle_id)
    with _LOCK:
        return apply_batch_cue_replacements(
            directory,
            payload.model_dump(),
            payload.replacements,
            actor,
        )


def apply_batch_cue_replacements(
    directory: Path, payload: dict[str, Any], replacements: list[dict[str, Any]], actor: str
) -> dict[str, Any]:
    from app.subtitles import canonical_state

    cues, source = canonical_state.canonical_cues(directory)
    if source != "ai_review_active":
        raise HTTPException(
            status_code=409,
            detail="目前無 AI Active Revision；batch replace 請使用既有 Editor 流程",
        )
    state = _review_state(directory)
    expected_sha = payload.get("base_content_sha256")
    identity = canonical_state.publication_identity(directory)
    if expected_sha is not None and expected_sha != identity["content_sha256"]:
        raise HTTPException(
            status_code=409, detail="Canonical content changed since load；請重新載入"
        )
    by_lineage = {tuple(str(v) for v in cue["source_segment_ids"]): dict(cue) for cue in cues}
    for item in replacements:
        if not isinstance(item, dict):
            raise HTTPException(status_code=422, detail="replacement 項目格式錯誤")
        lineage = tuple(str(v) for v in item.get("source_segment_ids", []))
        target = by_lineage.get(lineage)
        if target is None:
            raise HTTPException(status_code=404, detail=f"找不到 canonical cue: {list(lineage)}")
        if str(target.get("text")) != str(item.get("current_text")):
            raise HTTPException(status_code=409, detail=f"cue 文字已改變: {list(lineage)}")
        new_text = str(item.get("new_text", "")).strip()
        if not new_text:
            raise HTTPException(status_code=422, detail="new_text 不可空白")
        target["text"] = new_text
    new_cues = sorted(by_lineage.values(), key=lambda c: (c["start_ms"], c["end_ms"]))
    _validate_revision_cues(directory, new_cues)
    number = state["revision"] + 1
    digest = _canonical_cue_hash(new_cues)
    record = {
        "revision": number,
        "created_at": _iso(),
        "created_by": actor,
        "source": "editor_batch_replace",
        "edited_lineages": [list(t) for t in by_lineage],
        "accepted_change_ids": [],
        "rejected_change_ids": [],
        "cues": [
            {
                "text": str(cue["text"]),
                "source_segment_ids": [str(v) for v in cue["source_segment_ids"]],
                "start_ms": int(cue["start_ms"]),
                "end_ms": int(cue["end_ms"]),
            }
            for cue in new_cues
        ],
        "content_sha256": digest,
    }
    state["revisions"].append(record)
    state["active_revision"] = number
    state["revision"] = number
    _save_state(directory, state)
    return {
        "revision": number,
        "active_revision": number,
        "publication_key": f"ai_review_active:r{number}:{digest[:16]}",
        "replaced_count": len(replacements),
    }


@router.get("/{subtitle_id}/ai-review/baseline")
def get_baseline(subtitle_id: str) -> dict[str, Any]:
    """Immutable source evidence + canonical working representation.

    - source_segments: raw ASR evidence (never mutated).
    - working_cues: canonical current cues with lineage. With an Active AI
      revision these are the revision cues (merged/reflowed preserved as-is,
      never faked into per-segment text). Without one, they are the editor's
      corrected/manual per-segment state — never raw ASR.
    """
    directory = _directory(subtitle_id)
    baseline = baseline_segments(directory)
    from app.subtitles import canonical_state

    cues, canonical_source = canonical_state.canonical_cues(directory)
    return {
        "subtitle_id": subtitle_id,
        "canonical_source": canonical_source,
        "source_segments": baseline,
        "working_cues": [
            {
                "cue_id": cue.get("cue_id"),
                "text": cue["text"],
                "source_segment_ids": cue["source_segment_ids"],
                "start_ms": cue["start_ms"],
                "end_ms": cue["end_ms"],
            }
            for cue in cues
        ],
        # Legacy per-segment view: only meaningful when each segment maps to
        # its own cue; omitted for merged/reflowed states to avoid fake
        # per-segment working_text.
        **(
            {}
            if canonical_source == "ai_review_active" and any(
                len(cue["source_segment_ids"]) > 1 for cue in cues
            )
            else {
                "segments": [
                    {
                        **item,
                        "working_text": next(
                            (
                                cue["text"]
                                for cue in cues
                                if cue["source_segment_ids"] == [item["segment_id"]]
                            ),
                            item["raw_text"],
                        ),
                    }
                    for item in baseline
                ]
            }
        ),
    }
