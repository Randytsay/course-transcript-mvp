"""Create deterministic, local-only exports from immutable subtitle segments."""
from __future__ import annotations

import csv
import hashlib
import json
import os
from pathlib import Path

from docx import Document
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.styles import getSampleStyleSheet

DATA_DIR = Path(os.environ.get("COURSE_TRANSCRIPT_DATA_DIR", "/app/data"))
JOB = DATA_DIR / "jobs" / os.environ.get(
    "JOB_NAME", "voice_11386603-seg1"
)


def atomic_text(path: Path, text: str) -> None:
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(text, encoding="utf-8")
    temporary.replace(path)


def atomic_bytes(path: Path, value: bytes) -> None:
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_bytes(value)
    temporary.replace(path)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def timestamp(value: int) -> str:
    hours, value = divmod(max(0, value), 3_600_000)
    minutes, value = divmod(value, 60_000)
    seconds, _ = divmod(value, 1_000)
    return f"{hours:02}:{minutes:02}:{seconds:02}"


def ass_time(value: int) -> str:
    hours, value = divmod(max(0, value), 3_600_000)
    minutes, value = divmod(value, 60_000)
    seconds, milliseconds = divmod(value, 1_000)
    return f"{hours}:{minutes:02}:{seconds:02}.{milliseconds // 10:02}"


def ass_text(value: str) -> str:
    return (
        value.replace("\\", "\\\\")
        .replace("{", "\\{")
        .replace("}", "\\}")
        .replace("\n", "\\N")
    )


def _write_docx(path: Path, segments: list[dict], text_key: str) -> None:
    document = Document()
    document.add_heading("Course Transcript", level=0)
    document.add_paragraph(
        "時間碼來源：Google Cloud Speech-to-Text Chirp 3 word timestamps"
    )
    for segment in segments:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[{timestamp(int(segment['start_ms']))}] ").bold = True
        paragraph.add_run(str(segment[text_key]))
    temporary = path.with_suffix(path.suffix + ".tmp")
    document.save(temporary)
    temporary.replace(path)


def _pdf_font() -> str:
    # ReportLab's built-in CJK CID font avoids a platform-specific font file
    # dependency while preserving Traditional-Chinese text in the PDF.
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light"


def _write_pdf(path: Path, segments: list[dict], text_key: str) -> None:
    temporary = path.with_suffix(path.suffix + ".tmp")
    styles = getSampleStyleSheet()
    font = _pdf_font()
    title = styles["Title"]
    title.fontName = font
    body = styles["BodyText"]
    body.fontName = font
    body.fontSize = 10
    body.leading = 15
    story = [Paragraph("Course Transcript", title), Spacer(1, 12)]
    for segment in segments:
        value = (
            f"[{timestamp(int(segment['start_ms']))}] "
            f"{str(segment[text_key])}"
        )
        story.extend([Paragraph(value, body), Spacer(1, 5)])
    SimpleDocTemplate(str(temporary)).build(story)
    temporary.replace(path)


def _write_ass(path: Path, segments: list[dict], text_key: str) -> None:
    header = """[Script Info]
Title: Course Transcript MVP export
ScriptType: v4.00+
Collisions: Normal
PlayResX: 1920
PlayResY: 1080

[V4+ Styles]
Format: Name,Fontname,Fontsize,PrimaryColour,SecondaryColour,OutlineColour,BackColour,Bold,Italic,Underline,StrikeOut,ScaleX,ScaleY,Spacing,Angle,BorderStyle,Outline,Shadow,Alignment,MarginL,MarginR,MarginV,Encoding
Style: Default,Noto Sans CJK TC,48,&H00FFFFFF,&H000000FF,&H00101010,&H80000000,0,0,0,0,100,100,0,0,1,2,1,2,80,80,60,1

[Events]
Format: Layer,Start,End,Style,Name,MarginL,MarginR,MarginV,Effect,Text
"""
    lines = [header.rstrip("\n")]
    for segment in segments:
        lines.append(
            "Dialogue: 0,"
            f"{ass_time(int(segment['start_ms']))},"
            f"{ass_time(int(segment['end_ms']))},"
            "Default,,0,0,0,,"
            f"{ass_text(str(segment[text_key]))}"
        )
    atomic_text(path, "\n".join(lines) + "\n")


def main() -> int:
    raw = json.loads((JOB / "subtitles.json").read_text(encoding="utf-8"))
    corrected_path = JOB / "subtitles-corrected.json"
    corrected = (
        json.loads(corrected_path.read_text(encoding="utf-8"))
        if corrected_path.exists()
        else None
    )
    raw_segments = raw["segments"]
    corrected_by_id = (
        {item["segment_id"]: item for item in corrected["segments"]}
        if corrected
        else {}
    )
    published_segments = []
    for segment in raw_segments:
        correction = corrected_by_id.get(segment["segment_id"], {})
        published_segments.append(
            {
                **segment,
                "corrected_text": correction.get(
                    "corrected_text", segment["raw_text"]
                ),
                "uncertain_terms": correction.get("uncertain_terms", []),
            }
        )

    _write_ass(JOB / "subtitles.ass", raw_segments, "raw_text")
    if corrected:
        _write_ass(
            JOB / "subtitles-corrected.ass",
            published_segments,
            "corrected_text",
        )

    csv_path = JOB / "transcript.csv"
    with csv_path.open("w", encoding="utf-8", newline="") as stream:
        writer = csv.DictWriter(
            stream,
            fieldnames=[
                "segment_id",
                "start_ms",
                "end_ms",
                "raw_text",
                "corrected_text",
                "uncertain_terms",
            ],
        )
        writer.writeheader()
        for segment in published_segments:
            writer.writerow(
                {
                    **{
                        key: segment[key]
                        for key in ("segment_id", "start_ms", "end_ms", "raw_text")
                    },
                    "corrected_text": segment["corrected_text"],
                    "uncertain_terms": " | ".join(segment["uncertain_terms"]),
                }
            )
    atomic_bytes(JOB / "transcript-segments.csv", csv_path.read_bytes())

    transcript_json = {
        "job_id": JOB.name,
        "source_timing": "chirp_3 word timestamps",
        "correction_model": corrected.get("model") if corrected else None,
        "segments": published_segments,
    }
    atomic_text(
        JOB / "transcript.json",
        json.dumps(transcript_json, ensure_ascii=False, indent=2) + "\n",
    )
    atomic_text(
        JOB / "glossary_decisions.yaml",
        json.dumps(
            {"scope": "session", "decisions": [], "source": "human_review"},
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
    )
    join_path = JOB / "join-qa.json"
    if join_path.exists():
        atomic_bytes(JOB / "join_qa.json", join_path.read_bytes())
    glossary_source = JOB / "glossary" / "global-terms.csv"
    if glossary_source.exists():
        atomic_bytes(JOB / "glossary_candidates.csv", glossary_source.read_bytes())
    else:
        atomic_text(
            JOB / "glossary_candidates.csv",
            "canonical,variants,confidence\n",
        )

    # Canonical public artifact names. Legacy hyphenated names remain available
    # for compatibility with already completed jobs.
    atomic_bytes(JOB / "transcript_raw.txt", (JOB / "transcript-raw.txt").read_bytes())
    atomic_bytes(
        JOB / "transcript_timestamped.txt",
        (JOB / "transcript-timestamped.txt").read_bytes(),
    )
    corrected_text_path = JOB / "transcript-corrected.txt"
    atomic_bytes(
        JOB / "transcript_corrected.txt",
        (
            corrected_text_path
            if corrected_text_path.exists()
            else JOB / "transcript-raw.txt"
        ).read_bytes(),
    )
    canonical_srt = (
        JOB / "subtitles-corrected.srt"
        if (JOB / "subtitles-corrected.srt").exists()
        else JOB / "subtitles.srt"
    )
    canonical_vtt = (
        JOB / "subtitles-corrected.vtt"
        if (JOB / "subtitles-corrected.vtt").exists()
        else JOB / "subtitles.vtt"
    )
    atomic_bytes(JOB / "transcript.srt", canonical_srt.read_bytes())
    atomic_bytes(JOB / "transcript.vtt", canonical_vtt.read_bytes())

    _write_docx(JOB / "transcript.docx", published_segments, "corrected_text")
    _write_pdf(JOB / "transcript.pdf", published_segments, "corrected_text")

    generated = [
        "subtitles.srt",
        "subtitles.vtt",
        "subtitles.ass",
        "subtitles.json",
        "transcript-raw.txt",
        "transcript-timestamped.txt",
        "transcript.json",
        "transcript.csv",
        "transcript-segments.csv",
        "transcript.docx",
        "transcript.pdf",
        "transcript_raw.txt",
        "transcript_corrected.txt",
        "transcript_timestamped.txt",
        "transcript.srt",
        "transcript.vtt",
        "glossary_candidates.csv",
        "glossary_decisions.yaml",
        "join_qa.json",
    ]
    if corrected:
        generated.extend(
            [
                "subtitles-corrected.srt",
                "subtitles-corrected.vtt",
                "subtitles-corrected.ass",
                "subtitles-corrected.json",
                "transcript-corrected.txt",
            ]
        )
    artifacts = []
    for name in generated:
        path = JOB / name
        if not path.is_file() or path.stat().st_size <= 0:
            raise RuntimeError(f"missing or empty export: {name}")
        # Read back every output before recording it as a valid export.
        with path.open("rb") as stream:
            while stream.read(1024 * 1024):
                pass
        artifacts.append(
            {
                "name": name,
                "size_bytes": path.stat().st_size,
                "sha256": sha256(path),
                "source": "immutable Chirp segments",
            }
        )
    manifest = {
        "scope": "local_only",
        "drive_upload_started": False,
        "source_timing": "chirp_3 word timestamps",
        "correction_model": corrected.get("model") if corrected else None,
        "artifacts": artifacts,
    }
    atomic_text(
        JOB / "export-manifest.json",
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
    )
    print(
        f"EXPORT=PASS segments={len(raw_segments)} "
        f"corrected={bool(corrected)} artifacts={len(artifacts)}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
